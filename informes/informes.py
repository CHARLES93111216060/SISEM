import os
import sys
import logging
from docx import Document
from docx2pdf import convert
import json
import tempfile

# Allow running this module directly from the informes directory
PROJECT_ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
if PROJECT_ROOT not in sys.path:
    sys.path.insert(0, PROJECT_ROOT)

from database import db

logger = logging.getLogger(__name__)


def generar_reporte_pdf_concentrador(reporte_id):
    try:
        reporte = db.obtener_reporte_por_id(reporte_id)
        if not reporte:
            logger.error(f"No se encontró el reporte con ID {reporte_id}")
            return None

        plantilla_path = os.path.normpath(os.path.join(os.path.dirname(__file__), '..', 'Formatos', 'CONCENTRADOR_O2.docx'))
        if not os.path.exists(plantilla_path):
            logger.error(f"No se encontró el formato: {plantilla_path}")
            return None

        doc = Document(plantilla_path)
        reporte_raw = reporte[5]
        if not reporte_raw:
            logger.error("El campo del reporte está vacío")
            return None

        reporte_data = json.loads(reporte_raw)

        def lookup_nombre(func, clave):
            valor = reporte_data.get(clave)
            if not valor:
                return ''
            try:
                lista = func()
                indice = int(valor) - 1
                if 0 <= indice < len(lista):
                    return lista[indice][1]
            except Exception:
                pass
            return ''

        fecha = reporte_data.get('general.fecha_mantenimiento', '')
        equipo = lookup_nombre(db.tipo_equipos, 'equipo.tipo')
        cliente = lookup_nombre(db.clientes, 'cliente.nombre')
        area = reporte_data.get('cliente.area', '')
        marca = lookup_nombre(db.marcas, 'equipo.marca')
        modelo = reporte_data.get('equipo.modelo', '')
        serie = reporte_data.get('equipo.numero_serie', '')
        inventario = reporte_data.get('equipo.numero_inventario', '')
        servicio = lookup_nombre(db.servicios, 'servicio.tipo')
        limpieza = reporte_data.get('servicio.limpieza', '')
        pantalla = reporte_data.get('verificación de pantalla, botón de encendido.', '')
        llantas = reporte_data.get('verificación verificación de llantas de transporte.  ', '')
        flujometro = reporte_data.get('verificación de flujómetro.', '')
        cable = reporte_data.get('verificación de cable A.C.', '')
        bateria = reporte_data.get('verificación de bateria.', '')
        compresor = reporte_data.get('Horas funcionamiento de turbina y/o compresor.', '')
        filtro = reporte_data.get('Cambio de filtro de entrada de aire.', '')
        concentracion = reporte_data.get('Porcentaje de concentración de oxígeno.', '')
        flujo1 = reporte_data.get('Flujo 1', '')
        flujo2 = reporte_data.get('Flujo 2', '')
        flujo3 = reporte_data.get('Flujo 3', '')
        flujo4 = reporte_data.get('Flujo 4', '')
        flujo5 = reporte_data.get('Flujo 5', '')
        flujo6 = reporte_data.get('Flujo 6', '')
        flujo7 = reporte_data.get('Flujo 7', '')
        flujo8 = reporte_data.get('Flujo 8', '')
        flujo9 = reporte_data.get('Flujo 9', '')
        flujo10 = reporte_data.get('Flujo 10', '')
        sonora = reporte_data.get('Alarmas sonoras.', '')
        visual = reporte_data.get('Alarmas visuales (leds).', '')
        actividad = reporte_data.get('servicio.actividad', '')
        observacion = reporte_data.get('servicio.observaciones', '')
        cantidad = reporte_data.get('repuestos.cantidad_1', '')
        descripcion = reporte_data.get('repuestos.descripcion_1', '')
        estado = reporte_data.get('resultado.estado', '')
        profesional = lookup_nombre(db.personal, 'tecnico.nombre')
        report = reporte[0]

        reemplazos = {
            '<<Fecha del Mantenimiento:>>': fecha,
            'equipo': equipo,
            '<<Hosp. / Clínica:>>': cliente,
            '<<Área / Servicio:>>': area,
            '<<Marca:>>': marca,
            '<<Modelo:>>': modelo,
            '<<Serie:>>': serie,
            '<<N° inventario:>>': inventario,
            '<<TIPO DE SERVICIO REALIZADO.>>': servicio,
            '<<1>>': limpieza,
            '<<2>>':pantalla,
            '<<3>>': llantas,
            '<<4>>': flujometro,
            '<<5>>': cable,
            '<<6>>': bateria,
            '<<7>>': compresor,
            '<<8>>': filtro,
            '<<9>>': concentracion,
            '<<F_1>>': flujo1,
            '<<F_2>>': flujo2,
            '<<F_3>>': flujo3,
            '<<F_4>>': flujo4,
            '<<F_5>>': flujo5,
            '<<F_6>>': flujo6,
            '<<F_7>>': flujo7,
            '<<F_8>>': flujo8,
            '<<F_9>>': flujo9,
            '<<F_10>>': flujo10,
            '<<10>>': sonora,
            '<<11>>': visual,
            '.<<ACTIVIDAD REALIZADA.>>': actividad,
            '<<OBSERVACIONES.>>': observacion,
            '<<Cant_1.>>': cantidad,
            '<<Des_1.>>': descripcion,
            '<<ESTADO FINAL DE EQUIPO.>>': estado, 
            '<<ENTREGADO / REALIZADO POR:>>': profesional,
            '<<N° de Mant.>>': report
        }

        for p in doc.paragraphs:
            for key, value in reemplazos.items():
                if key in p.text:
                    p.text = p.text.replace(key, str(value))

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for key, value in reemplazos.items():
                        if key in cell.text:
                            cell.text = cell.text.replace(key, str(value))

        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as temporal:
            temporal_path = temporal.name
            doc.save(temporal_path)

        output_pdf = temporal_path.replace('.docx', '.pdf')
        co_inited = False
        try:
            # Inicializar COM en Windows para evitar 'CoInitialize has not been called.'
            try:
                import pythoncom
                pythoncom.CoInitialize()
                co_inited = True
            except Exception:
                co_inited = False

            convert(temporal_path, output_pdf)
            if os.path.exists(output_pdf):
                return output_pdf
            return None

        except Exception as e:
            logger.exception(f"Error al convertir a PDF: {e}")
            raise

        finally:
            if co_inited:
                try:
                    pythoncom.CoUninitialize()
                except Exception:
                    pass
            if os.path.exists(temporal_path):
                os.remove(temporal_path)

    except Exception as e:
        logger.exception(f"Error al generar el reporte: {e}")
        # Propagar para que el caller (app) pueda devolver el error completo
        raise


